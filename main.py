import streamlit as st
import tempfile
import os
from pathlib import Path
import io
from PIL import Image
import fitz  # PyMuPDF
from docx import Document
import easyocr
import numpy as np
import zipfile

# Translation dictionary
TRANSLATIONS = {
    "en": {
        "title": "📄 PDF OCR to DOCX Converter",
        "description": "Upload one or multiple PDF files to extract text using OCR and convert to DOCX format.",
        "language_label": "Language / Sprache",
        "loading_model": "Loading OCR model... (This may take a moment on first run)",
        "model_load_error": "❌ Failed to load OCR reader. Please refresh the page and try again.",
        "model_loaded": "✅ OCR model loaded successfully!",
        "instructions_title": "ℹ️ Instructions",
        "instructions_content": """
        1. **Upload PDFs**: Drag and drop or browse to select one or multiple PDF files
        2. **Processing**: The app will convert each PDF page to images and perform OCR using EasyOCR
        3. **Download**: Download the generated DOCX files with extracted text

        **Features**:
        - ✅ Works on Streamlit Cloud (no system dependencies required)
        - ✅ Supports German and English text recognition
        - ✅ High-quality OCR using deep learning
        - ✅ Batch processing of multiple files
        - ✅ Text preview before download

        **Tips for better results**:
        - Use high-quality, clear PDF files
        - Ensure text is not too small or blurry
        - Processing time depends on the number of pages
        """,
        "file_uploader_label": "Choose PDF files",
        "file_uploader_help": "You can upload multiple PDF files at once",
        "files_uploaded": "file(s) uploaded",
        "start_processing": "🚀 Start OCR Processing",
        "processing_file": "Processing",
        "converting_pdf": "Converting PDF to images...",
        "pages_converted": "pages to images",
        "performing_ocr": "Performing OCR...",
        "processing_page": "Processing page",
        "no_text_extracted": "⚠️ No text was extracted from this PDF",
        "creating_docx": "Creating DOCX document...",
        "download_button": "📥 Download",
        "preview_text": "📖 Preview extracted text from",
        "extracted_text_label": "Extracted text:",
        "successfully_processed": "Successfully processed",
        "failed_to_process": "Failed to process",
        "processing_complete": "🎉 Processing complete!",
        "files_processed": "files processed successfully.",
        "no_files_processed": "❌ No files were processed successfully. Please check your PDF files and try again.",
        "footer_technology": "**Technology**: This app uses PyMuPDF for PDF processing and EasyOCR for text extraction.\nBoth libraries work great on Streamlit Cloud without requiring system dependencies!",
        "footer_languages": "**Supported Languages**: German and English are automatically detected and processed.",
        "page_separator": "Page"
    },
    "de": {
        "title": "📄 PDF OCR zu DOCX Konverter",
        "description": "Laden Sie eine oder mehrere PDF-Dateien hoch, um Text mittels OCR zu extrahieren und in das DOCX-Format zu konvertieren.",
        "language_label": "Sprache / Language",
        "loading_model": "OCR-Modell wird geladen... (Dies kann beim ersten Start einen Moment dauern)",
        "model_load_error": "❌ OCR-Modell konnte nicht geladen werden. Bitte aktualisieren Sie die Seite und versuchen Sie es erneut.",
        "model_loaded": "✅ OCR-Modell erfolgreich geladen!",
        "instructions_title": "ℹ️ Anleitung",
        "instructions_content": """
        1. **PDFs hochladen**: Ziehen Sie PDF-Dateien per Drag & Drop hierher oder wählen Sie eine oder mehrere Dateien aus
        2. **Verarbeitung**: Die App konvertiert jede PDF-Seite in Bilder und führt OCR mit EasyOCR durch
        3. **Download**: Laden Sie die generierten DOCX-Dateien mit dem extrahierten Text herunter

        **Funktionen**:
        - ✅ Funktioniert auf Streamlit Cloud (keine Systemabhängigkeiten erforderlich)
        - ✅ Unterstützt deutsche und englische Texterkennung
        - ✅ Hochwertige OCR mit Deep Learning
        - ✅ Stapelverarbeitung mehrerer Dateien
        - ✅ Textvorschau vor dem Download

        **Tipps für bessere Ergebnisse**:
        - Verwenden Sie hochwertige, klare PDF-Dateien
        - Stellen Sie sicher, dass der Text nicht zu klein oder unscharf ist
        - Die Verarbeitungszeit hängt von der Anzahl der Seiten ab
        """,
        "file_uploader_label": "PDF-Dateien auswählen",
        "file_uploader_help": "Sie können mehrere PDF-Dateien gleichzeitig hochladen",
        "files_uploaded": "Datei(en) hochgeladen",
        "start_processing": "🚀 OCR-Verarbeitung starten",
        "processing_file": "Verarbeitung",
        "converting_pdf": "PDF wird zu Bildern konvertiert...",
        "pages_converted": "Seiten zu Bildern konvertiert",
        "performing_ocr": "OCR wird durchgeführt...",
        "processing_page": "Verarbeite Seite",
        "no_text_extracted": "⚠️ Kein Text wurde aus dieser PDF extrahiert",
        "creating_docx": "DOCX-Dokument wird erstellt...",
        "download_button": "📥 Download",
        "preview_text": "📖 Vorschau des extrahierten Texts aus",
        "extracted_text_label": "Extrahierter Text:",
        "successfully_processed": "erfolgreich verarbeitet",
        "failed_to_process": "Fehler beim Verarbeiten von",
        "processing_complete": "🎉 Verarbeitung abgeschlossen!",
        "files_processed": "Dateien erfolgreich verarbeitet.",
        "no_files_processed": "❌ Keine Dateien wurden erfolgreich verarbeitet. Bitte überprüfen Sie Ihre PDF-Dateien und versuchen Sie es erneut.",
        "footer_technology": "**Technologie**: Diese App verwendet PyMuPDF für PDF-Verarbeitung und EasyOCR für Texterkennung.\nBeide Bibliotheken funktionieren hervorragend auf Streamlit Cloud ohne Systemabhängigkeiten!",
        "footer_languages": "**Unterstützte Sprachen**: Deutsch und Englisch werden automatisch erkannt und verarbeitet.",
        "page_separator": "Seite"
    }
}


def setup_page():
    """Configure the Streamlit page"""
    st.set_page_config(
        page_title="PDF OCR to DOCX Converter",
        page_icon="📄",
        layout="wide"
    )


def get_text(key, lang="en"):
    """Get translated text"""
    return TRANSLATIONS.get(lang, TRANSLATIONS["en"]).get(key, key)


@st.cache_resource
def load_ocr_reader():
    """Load EasyOCR reader (cached to avoid reloading)"""
    try:
        reader = easyocr.Reader(['en', 'de'])  # English and German support
        return reader
    except Exception as e:
        st.error(f"Error loading OCR reader: {str(e)}")
        return None


def pdf_to_images(pdf_bytes):
    """Convert PDF pages to images using PyMuPDF"""
    try:
        images = []
        # Open PDF from bytes
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")

        for page_num in range(pdf_document.page_count):
            # Get page
            page = pdf_document.load_page(page_num)

            # Convert to image (higher resolution for better OCR)
            mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better quality
            pix = page.get_pixmap(matrix=mat)

            # Convert to PIL Image
            img_data = pix.tobytes("ppm")
            img = Image.open(io.BytesIO(img_data))
            images.append(img)

        pdf_document.close()
        return images

    except Exception as e:
        st.error(f"Error converting PDF to images: {str(e)}")
        return None


def perform_ocr_easyocr(images, reader, lang="en"):
    """Perform OCR using EasyOCR"""
    extracted_text = ""

    progress_bar = st.progress(0)
    status_text = st.empty()

    for i, image in enumerate(images):
        try:
            # Update progress
            progress = (i + 1) / len(images)
            progress_bar.progress(progress)

            # Fix: Use simple string formatting instead of get_text for "of"
            of_text = "von" if lang == 'de' else "of"
            status_text.text(
                f"{get_text('processing_page', lang)} {i + 1} {of_text} {len(images)}...")

            # Convert PIL image to numpy array
            img_array = np.array(image)

            # Perform OCR
            results = reader.readtext(img_array, paragraph=True)

            # Extract text from results
            page_text = ""
            for detection in results:
                text = detection[1]  # detection[0] is bbox, detection[1] is text, detection[2] is confidence
                page_text += text + " "

            extracted_text += f"\n--- {get_text('page_separator', lang)} {i + 1} ---\n"
            extracted_text += page_text.strip()
            extracted_text += "\n\n"

        except Exception as e:
            st.warning(f"Error performing OCR on page {i + 1}: {str(e)}")
            continue

    progress_bar.empty()
    status_text.empty()

    return extracted_text


def create_docx(text, filename, lang="en"):
    """Create a DOCX document from extracted text"""
    try:
        doc = Document()

        # Add title
        title = doc.add_heading(f'OCR Result: {filename}', 0)

        # Add the extracted text
        # Split text into paragraphs for better formatting
        paragraphs = text.split('\n\n')

        for paragraph in paragraphs:
            if paragraph.strip():
                # Handle page separators
                if paragraph.strip().startswith(f'--- {get_text("page_separator", lang)}'):
                    doc.add_heading(paragraph.strip(), level=1)
                else:
                    doc.add_paragraph(paragraph.strip())

        return doc
    except Exception as e:
        st.error(f"Error creating DOCX: {str(e)}")
        return None


def create_download_link(doc, filename, lang="en"):
    """Create a download link for the DOCX document"""
    try:
        # Save document to bytes
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        # Create download button
        st.download_button(
            label=f"{get_text('download_button', lang)} {filename}",
            data=doc_io.getvalue(),
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"download_{filename}"
        )

        return True
    except Exception as e:
        st.error(f"Error creating download link: {str(e)}")
        return False


def process_single_pdf(uploaded_file, reader, lang="en"):
    """Process a single PDF file"""
    try:
        # Get PDF bytes
        pdf_bytes = uploaded_file.getvalue()

        # Convert PDF to images
        with st.spinner(get_text("converting_pdf", lang)):
            images = pdf_to_images(pdf_bytes)

        if images is None:
            return False

        st.success(f"✅ {len(images)} {get_text('pages_converted', lang)}")

        # Perform OCR
        with st.spinner(get_text("performing_ocr", lang)):
            extracted_text = perform_ocr_easyocr(images, reader, lang)

        if not extracted_text.strip():
            st.warning(get_text("no_text_extracted", lang))
            return False

        # Create DOCX
        with st.spinner(get_text("creating_docx", lang)):
            original_name = Path(uploaded_file.name).stem
            docx_filename = f"{original_name}_OCR.docx"
            doc = create_docx(extracted_text, original_name, lang)

        if doc is None:
            return False

        # Create download link
        create_download_link(doc, docx_filename, lang)

        # Show preview of extracted text
        with st.expander(f"{get_text('preview_text', lang)} {uploaded_file.name}"):
            st.text_area(
                get_text("extracted_text_label", lang),
                extracted_text[:2000] + ("..." if len(extracted_text) > 2000 else ""),
                height=200,
                key=f"preview_{uploaded_file.name}_{hash(extracted_text[:100])}"
            )

        return True

    except Exception as e:
        st.error(f"Error processing {uploaded_file.name}: {str(e)}")
        return False


def create_zip_download(processed_files, lang="en"):
    """Create a ZIP file containing all DOCX files for download"""
    try:
        zip_buffer = io.BytesIO()

        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for filename, doc_bytes in processed_files.items():
                zip_file.writestr(filename, doc_bytes)

        zip_buffer.seek(0)

        # Create download button for ZIP
        st.download_button(
            label=f"📦 Download All Files ({len(processed_files)} files)",
            data=zip_buffer.getvalue(),
            file_name="ocr_results.zip",
            mime="application/zip",
            key="download_all_zip"
        )

        return True
    except Exception as e:
        st.error(f"Error creating ZIP file: {str(e)}")
        return False


def process_single_pdf_with_state(uploaded_file, reader, lang="en"):
    """Process a single PDF file and store result in session state"""
    try:
        # Get PDF bytes
        pdf_bytes = uploaded_file.getvalue()

        # Convert PDF to images
        with st.spinner(get_text("converting_pdf", lang)):
            images = pdf_to_images(pdf_bytes)

        if images is None:
            return False

        st.success(f"✅ {len(images)} {get_text('pages_converted', lang)}")

        # Perform OCR
        with st.spinner(get_text("performing_ocr", lang)):
            extracted_text = perform_ocr_easyocr(images, reader, lang)

        if not extracted_text.strip():
            st.warning(get_text("no_text_extracted", lang))
            return False

        # Create DOCX
        with st.spinner(get_text("creating_docx", lang)):
            original_name = Path(uploaded_file.name).stem
            docx_filename = f"{original_name}_OCR.docx"
            doc = create_docx(extracted_text, original_name, lang)

        if doc is None:
            return False

        # Save document to bytes and store in session state
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_bytes = doc_io.getvalue()

        # Store in session state
        if 'processed_files' not in st.session_state:
            st.session_state.processed_files = {}

        st.session_state.processed_files[docx_filename] = {
            'doc_bytes': doc_bytes,
            'extracted_text': extracted_text,
            'original_name': uploaded_file.name
        }

        # Show preview of extracted text
        with st.expander(f"{get_text('preview_text', lang)} {uploaded_file.name}"):
            st.text_area(
                get_text("extracted_text_label", lang),
                extracted_text[:2000] + ("..." if len(extracted_text) > 2000 else ""),
                height=200,
                key=f"preview_{uploaded_file.name}_{hash(extracted_text[:100])}"
            )

        return True

    except Exception as e:
        st.error(f"Error processing {uploaded_file.name}: {str(e)}")
        return False


def main():
    """Main application function"""
    setup_page()

    # Initialize session state
    if 'processed_files' not in st.session_state:
        st.session_state.processed_files = {}

    # Language selector in sidebar
    with st.sidebar:
        st.markdown("### Settings")
        lang = st.selectbox(
            "🌐 Language / Sprache",
            options=["en", "de"],
            format_func=lambda x: "🇺🇸 English" if x == "en" else "🇩🇪 Deutsch",
            index=0
        )

    # Header
    st.title(get_text("title", lang))
    st.markdown(get_text("description", lang))

    # Load OCR reader
    with st.spinner(get_text("loading_model", lang)):
        reader = load_ocr_reader()

    if reader is None:
        st.error(get_text("model_load_error", lang))
        return

    st.success(get_text("model_loaded", lang))

    # Instructions
    with st.expander(get_text("instructions_title", lang)):
        st.markdown(get_text("instructions_content", lang))

    # File uploader
    uploaded_files = st.file_uploader(
        get_text("file_uploader_label", lang),
        type=['pdf'],
        accept_multiple_files=True,
        help=get_text("file_uploader_help", lang)
    )

    if uploaded_files:
        st.info(f"📁 {len(uploaded_files)} {get_text('files_uploaded', lang)}")

        # Clear previous results button
        if st.session_state.processed_files:
            if st.button("🗑️ Clear Previous Results"):
                st.session_state.processed_files = {}
                st.rerun()

        # Process button
        if st.button(get_text("start_processing", lang), type="primary"):
            success_count = 0
            total_files = len(uploaded_files)

            # Process each file
            for i, uploaded_file in enumerate(uploaded_files):
                st.subheader(f"{get_text('processing_file', lang)} {i + 1}/{total_files}: {uploaded_file.name}")

                if process_single_pdf_with_state(uploaded_file, reader, lang):
                    success_count += 1
                    st.success(f"✅ {uploaded_file.name} {get_text('successfully_processed', lang)}")
                else:
                    st.error(f"❌ {get_text('failed_to_process', lang)} {uploaded_file.name}")

                st.divider()

            # Summary
            if success_count > 0:
                st.balloons()
                st.success(
                    f"{get_text('processing_complete', lang)} {success_count}/{total_files} {get_text('files_processed', lang)}")
            else:
                st.error(get_text("no_files_processed", lang))

    # Display processed files and download options
    if st.session_state.processed_files:
        st.markdown("---")
        st.subheader("📥 Download Results")

        # Download all as ZIP
        if len(st.session_state.processed_files) > 1:
            files_for_zip = {filename: data['doc_bytes'] for filename, data in st.session_state.processed_files.items()}
            create_zip_download(files_for_zip, lang)
            st.markdown("**Or download individual files:**")

        # Individual download buttons
        cols = st.columns(min(3, len(st.session_state.processed_files)))

        for idx, (filename, data) in enumerate(st.session_state.processed_files.items()):
            col = cols[idx % len(cols)]
            with col:
                st.download_button(
                    label=f"📥 {filename}",
                    data=data['doc_bytes'],
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_individual_{filename}"
                )

    # Footer
    st.markdown("---")
    st.markdown(get_text("footer_technology", lang))
    st.markdown(get_text("footer_languages", lang))


if __name__ == "__main__":
    main()
